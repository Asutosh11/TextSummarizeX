#!/usr/bin/env python
# -*- coding: utf-8 -*- 
from docx import Document

from io import StringIO

from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'dox'}


def readTxt(file):
    return file.read()

def readDocx(file):
    doc = Document(file)
    txt = ""
    for para in doc.paragraphs:
        txt = txt + para.text
    return txt

def readPdf(file):
    output_string = StringIO()
    parser = PDFParser(file)
    doc = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)

    return(output_string.getvalue())